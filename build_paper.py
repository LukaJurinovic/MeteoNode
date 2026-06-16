# -*- coding: utf-8 -*-
"""Generates the Croatian final-project paper as a .docx file.
Code snippets are read directly from the real source tree so they stay authentic.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))

def src(rel):
    with open(os.path.join(ROOT, rel), "r", encoding="utf-8") as f:
        return f.read().splitlines()

def lines(rel, start, end):
    """1-based inclusive line range from a source file."""
    return "\n".join(src(rel)[start-1:end])

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def shade(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexcolor)
    pPr.append(sh)

def code(text, caption=None):
    if caption:
        cp = doc.add_paragraph()
        r = cp.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        cp.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(6)
    shade(p, "F2F2F2")
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    # ensure east-asian/other fallback also Consolas
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    return p

def para(text):
    return doc.add_paragraph(text)

def h1(t):
    doc.add_heading(t, level=1)

def h2(t):
    doc.add_heading(t, level=2)

def h3(t):
    doc.add_heading(t, level=3)

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

def pagebreak():
    doc.add_page_break()

# =========================================================
# TITLE PAGE
# =========================================================
def center(run_text, size, bold=False, space=6, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(run_text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

for _ in range(2):
    doc.add_paragraph()
center("ZAVRŠNI RAD", 14, bold=True, space=4, color=RGBColor(0x55,0x55,0x55))
doc.add_paragraph()
center("MeteoNode", 30, bold=True, space=4, color=RGBColor(0x1F,0x3A,0x5F))
center("Platforma za nadzor okolišnih uvjeta temeljena na", 14, space=2)
center("senzorskim čvorovima bez WiFi povezivosti (SMS / GSM)", 14, space=2)
for _ in range(3):
    doc.add_paragraph()
center("Distribuirani IoT sustav: ESP32 senzorski čvorovi · SMS gateway ·", 12, space=2)
center("Spring Boot REST poslužitelj · React nadzorna ploča", 12, space=2)
for _ in range(6):
    doc.add_paragraph()
center("Autor: Luka Jurinović", 12, space=2)
center("Mentor: ___________________", 12, space=2)
center("Akademska godina 2025./2026.", 12, space=2)
pagebreak()

# =========================================================
# SAŽETAK
# =========================================================
h1("Sažetak")
para(
 "MeteoNode je samostalno hostana platforma za nadzor okolišnih uvjeta čija je "
 "ključna projektna pretpostavka da senzorski čvorovi nemaju pristup WiFi mreži. "
 "Čvorovi komuniciraju isključivo putem SMS poruka koristeći SIM800L GSM modul. "
 "Jedan ESP32 gateway premošćuje GSM mrežu i internet: prima SMS poruke od čvorova, "
 "prevodi ih u HTTP zahtjeve prema Spring Boot poslužitelju, te povratno isporučuje "
 "naredbe poslužitelja prema čvorovima ponovno putem SMS-a. React aplikacija pruža "
 "operaterima pregled svih postaja u stvarnom vremenu, povijesne grafove, alarmna "
 "pravila i daljinsko upravljanje čvorovima."
)
para(
 "Rad opisuje cjelokupnu arhitekturu sustava kroz četiri neovisno isporučiva "
 "artefakta — firmware senzorskog čvora, firmware gatewaya, backend i frontend — "
 "te detaljno analizira tri ključna procesna toka (engl. pipeline): tok mjerenja od "
 "senzora do baze, tok naredbi od poslužitelja do čvora, te tok evaluacije alarma. "
 "Poseban naglasak stavljen je na granicu protokola: SMS format poruka koji "
 "predstavlja ugovor između firmwarea i poslužitelja."
)
para("Ključne riječi: IoT, ESP32, GSM, SMS, Spring Boot, React, senzorske mreže, "
     "meteorološke postaje, distribuirani sustavi.")

h2("Abstract")
para(
 "MeteoNode is a self-hosted environmental monitoring platform built around a single "
 "hard constraint: sensor nodes have no WiFi connectivity. Nodes communicate "
 "exclusively over SMS using a SIM800L cellular module. A single ESP32 gateway bridges "
 "the cellular network and the internet, forwarding node data to a Spring Boot REST API "
 "over HTTP and relaying backend-issued commands back to nodes over SMS. A React "
 "dashboard provides operators with a live view of every station, historical charts, "
 "alarm rules and remote node control. This paper documents the full architecture and "
 "analyses three core pipelines: measurement ingestion, command delivery and alarm "
 "evaluation."
)
pagebreak()

# =========================================================
# SADRŽAJ (TOC field)
# =========================================================
h1("Sadržaj")
tp = doc.add_paragraph()
run = tp.add_run()
fldStart = OxmlElement("w:fldChar"); fldStart.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-3" \\h \\z \\u'
fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
hint = OxmlElement("w:t"); hint.text = "Desnim klikom → Update Field za generiranje sadržaja."
fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
for el in (fldStart, instr, fldSep, hint, fldEnd):
    run._r.append(el)
pagebreak()

# =========================================================
# 1. UVOD
# =========================================================
h1("1. Uvod")
h2("1.1. Motivacija")
para(
 "Tipične platforme za nadzor okoliša pretpostavljaju da senzorski uređaji imaju "
 "pouzdanu mrežnu vezu — najčešće WiFi ili Ethernet. Ta pretpostavka pada čim se "
 "uređaj postavi na udaljenu lokaciju: poljoprivredno zemljište, planinarski dom, "
 "mjerno mjesto uz prometnicu ili bilo koju točku izvan dosega bežične mreže. Na "
 "takvim lokacijama GSM signal je često jedina dostupna komunikacijska "
 "infrastruktura."
)
para(
 "MeteoNode polazi od te stvarnosti i pretvara ograničenje u temeljnu projektnu "
 "odluku: senzorski čvor nema IP adresu, nema WebSocket vezu, nema MQTT broker. "
 "Jedino što čvor može jest poslati SMS poruku i primiti SMS poruku. Cijela "
 "arhitektura sustava izgrađena je oko tog ograničenja."
)

h2("1.2. Pregled rješenja")
para(
 "Sustav se sastoji od četiri neovisno isporučiva dijela koji zajedno čine "
 "cjeloviti lanac od fizičkog senzora do operaterskog preglednika:"
)
bullets([
 "Senzorski čvor (ESP32 + SIM800L) — mjeri temperaturu, tlak, vlažnost, UV indeks i "
 "osvjetljenje te šalje rezultate SMS porukom.",
 "Gateway (ESP32 + SIM800L + WiFi) — jedini most između GSM mreže i interneta; "
 "prima SMS-ove, prosljeđuje mjerenja HTTP-om i isporučuje naredbe natrag čvorovima.",
 "Backend (Spring Boot 3 + MySQL) — REST API sa strogom četveroslojnom arhitekturom, "
 "JWT i API-ključ autentikacijom, alarmnim motorom i redom naredbi.",
 "Frontend (React + Vite + TypeScript) — nadzorna ploča s karticama postaja, "
 "povijesnim grafovima, alarmnim pravilima i upravljanjem uređajima.",
])

h2("1.3. Ciljevi rada")
para(
 "Cilj ovog rada jest dokumentirati cjelovitu arhitekturu MeteoNode sustava i "
 "objasniti kako pojedini dijelovi surađuju unatoč tome što su razvijani i "
 "isporučivani neovisno. Posebna pažnja posvećena je analizi triju procesnih tokova "
 "koji predstavljaju srž sustava — toku mjerenja, toku naredbi i toku alarma — jer "
 "upravo oni prelaze granice između firmwarea, GSM mreže, poslužitelja i baze."
)

h2("1.4. Struktura rada")
para(
 "Nakon uvoda, drugo poglavlje daje pregled arhitekture i relacijsku sliku "
 "komponenti. Treće poglavlje opisuje hardver. Četvrto, peto i šesto poglavlje "
 "posvećena su firmwareu — čvoru, gatewayu i zajedničkoj SMS biblioteci. Sedmo i "
 "osmo poglavlje opisuju backend i bazu podataka, deveto autentikaciju. Poglavlja "
 "deset, jedanaest i dvanaest detaljno razrađuju tri ključna procesna toka. Trinaesto "
 "poglavlje opisuje dinamičku registraciju senzora, četrnaesto frontend, a petnaesto "
 "testiranje. Rad završava poglavljima o pokretanju sustava i zaključkom."
)
pagebreak()

# =========================================================
# 2. ARHITEKTURA
# =========================================================
h1("2. Pregled sustava i arhitektura")
h2("2.1. Relacijska slika")
para("Komponente sustava i kanali kojima komuniciraju prikazani su sljedećom shemom:")
code(
 "[Senzorski cvor]  --SMS-->  [Gateway]  --HTTP-->  [Backend]  <--HTTP--  [React frontend]\n"
 "  ESP32 + SIM800L             ESP32                Spring Boot            korisnik u browseru\n"
 "  BMP280, DHT11               WiFi + SIM           MySQL baza\n"
 "  GUVA, BH1750",
 caption="Slika 2.1 — Komunikacijski kanali između komponenti"
)
para(
 "Ključna točka cijele arhitekture jest da čvor nikada ne komunicira s backendom "
 "izravno. Gateway je jedini most između GSM mreže i interneta. Čvor zna samo "
 "telefonski broj gatewaya; gateway zna samo URL backenda i svoj API ključ."
)

h2("2.2. Neovisno isporučivi artefakti")
para(
 "Sustav se sastoji od tri neovisno isporučiva artefakta: firmware čvora (flešanje "
 "ESP32 uređaja), firmware gatewaya (zaseban flešani uređaj) i backend (deploy "
 "poslužitelja). Granica koja ih povezuje jest SMS format poruka. Svaka promjena tog "
 "formata mora se koordinirano provesti kroz sva tri artefakta, bez mogućnosti "
 "jednostranog vraćanja unatrag, što je glavni izvor složenosti sustava."
)

h2("2.3. Procjena složenosti")
para("Okvirna veličina pojedinih slojeva sustava prikazana je u tablici:")
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "Sloj"; hdr[1].text = "Jezik"; hdr[2].text = "Približan broj linija"
for r in [("Backend", "Java (Spring Boot)", "~3.300"),
          ("Frontend", "TypeScript / TSX", "~2.100"),
          ("Firmware", "C++ / Arduino", "~1.100")]:
    cells = tbl.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = r
para(
 "Backend je sam po sebi rutinski — četveroslojni Spring Boot s JWT-om i "
 "autorizacijom. Složenost dolazi iz protokolarnog ugovora između slojeva: SMS "
 "format je granica između firmwarea i backenda, a stateful SMS protokol "
 "(prikvačene potvrde, grupirane naredbe, implicitni ACK, NVS perzistencija na obje "
 "strane) glavni je sekundarni izvor složenosti."
)
pagebreak()

# =========================================================
# 3. HARDVER
# =========================================================
h1("3. Hardver")
h2("3.1. Senzorski čvor")
para(
 "Senzorski čvor temelji se na ESP32 razvojnoj pločici na koju su spojena četiri "
 "senzora i jedan GSM modul. Pregled senzora i njihovih sučelja:"
)
bullets([
 "BMP280 — temperatura i tlak, I²C sabirnica na adresi 0x76 (s rezervom na 0x77).",
 "DHT11 — relativna vlažnost, digitalni pin GPIO 26.",
 "GUVA-S12SD — ultraljubičasto zračenje, analogni ulaz ADC na GPIO 34.",
 "GY-302 / BH1750 — ambijentalno osvjetljenje, I²C sabirnica na adresi 0x23.",
 "SIM800L — GSM modul, UART2 (RX=16, TX=17).",
])
para(
 "BMP280 i GY-302 dijele istu I²C sabirnicu (SDA=21, SCL=22). Definicije pinova u "
 "firmwareu čvora odgovaraju ovom rasporedu:"
)
code(lines("firmware/node/node.ino", 78, 88),
     caption="Listing 3.1 — firmware/node/node.ino: raspored pinova i konstante")

h2("3.2. Gateway")
para(
 "Gateway je također ESP32, ali za razliku od čvora posjeduje i WiFi i SIM800L. "
 "Njegova jedina zadaća jest prevođenje između dvaju svjetova: prima SMS poruke s "
 "GSM strane i šalje HTTP zahtjeve s WiFi strane. Gateway u memoriji drži registar "
 "do pet poznatih čvorova (serijski broj → ID čvora + ID-evi senzora) koji "
 "perzistira u SPIFFS flash datotečni sustav kako bi preživio restart."
)

h2("3.3. Otpornost senzorskog očitanja")
para(
 "Hardverska raznolikost donosi i specifične probleme robusnosti. BMP280 inicijalizacija "
 "automatski pokušava obje moguće I²C adrese i provjerava chip ID prije nego što "
 "proglasi senzor prisutnim:"
)
code(lines("firmware/node/node.ino", 26, 46),
     caption="Listing 3.2 — Inicijalizacija BMP280 senzora s detekcijom adrese")
para(
 "Analogni UV senzor nema mehanizam detekcije kvara, pa se njegova vrijednost "
 "ograničava (engl. sanity clamp): očitanja izvan razumnog raspona pretvaraju se u "
 "NaN i kasnije se ne prosljeđuju. Isto vrijedi i za negativne vrijednosti svjetla:"
)
code(lines("firmware/node/node.ino", 273, 289),
     caption="Listing 3.3 — Čitanje senzora s pretvaranjem nevažećih vrijednosti u NaN")
pagebreak()

# =========================================================
# 4. FIRMWARE - NODE
# =========================================================
h1("4. Firmware — senzorski čvor")
h2("4.1. Životni ciklus")
para(
 "Firmware čvora inicijalizira senzore pri pokretanju, učitava interval izvještavanja "
 "iz NVS memorije te ulazi u petlju: spava → čita → gradi SMS → šalje. Obrada naredbi "
 "odvija se sinkrono između ciklusa. Pri pokretanju se izvršava setup() funkcija koja "
 "konfigurira watchdog timer, učitava perzistirano stanje i šalje prvo očitanje:"
)
code(lines("firmware/node/node.ino", 116, 183),
     caption="Listing 4.1 — firmware/node/node.ino: setup() inicijalizacija čvora")

h2("4.2. Glavna petlja")
para(
 "Glavna petlja čvora ima dvije odgovornosti: osluškuje dolazne SMS događaje s GSM "
 "modula (prepoznaje +CMT: zaglavlje koje najavljuje tijelo poruke) i provjerava je "
 "li proteklo dovoljno vremena za sljedeće slanje očitanja. Watchdog se resetira na "
 "svakom prolazu kako bi se izbjegao panic restart."
)
code(lines("firmware/node/node.ino", 185, 218),
     caption="Listing 4.2 — Glavna petlja: prijem SMS-a i periodičko slanje")

h2("4.3. Perzistencija intervala u NVS")
para(
 "Interval izvještavanja (zadano 86400 s, tj. 24 sata) čuva se u NVS (Non-Volatile "
 "Storage) memoriji i preživljava prekid napajanja. Pri pokretanju se uz interval "
 "provjerava i identifikator implementacije (deployId); ako se promijenio, stanje se "
 "briše. Posebna logika postoji za migraciju starih vrijednosti koje su bile "
 "pohranjene u milisekundama umjesto u sekundama (vidi listing 4.1, linije 140–149)."
)

h2("4.4. Obrada naredbi i deduplikacija")
para(
 "Kada čvor primi naredbu, ona prolazi kroz funkciju handleCommand koja razlaže "
 "grupiranu poruku na pojedinačne naredbe i provjerava nije li svaka već nedavno "
 "obrađena. Deduplikacija po identifikatoru naredbe (unutar 120 sekundi) sprječava "
 "dvostruko izvršavanje ako GSM mreža ponovno isporuči istu poruku:"
)
code(lines("firmware/node/node.ino", 313, 341),
     caption="Listing 4.3 — Deduplikacija i obrada grupirane naredbe")
para("Pojedinačna naredba se zatim izvršava prema svom tipu:")
code(lines("firmware/node/node.ino", 291, 311),
     caption="Listing 4.4 — processCommand(): REBOOT, REQUEST_READINGS, SET_INTERVAL")
para(
 "Naredba REBOOT šalje potvrdu prije ponovnog pokretanja jer nakon ESP.restart() više "
 "ne bi imala priliku potvrditi izvršenje. Naredbe SET_INTERVAL i REQUEST_READINGS "
 "potvrdu prikvače na sljedeći podatkovni SMS putem polja pendingAcks."
)

h2("4.5. Slanje SMS poruke")
para(
 "Slanje SMS-a koristi standardne AT naredbe SIM800L modula i ima ugrađen mehanizam "
 "ponovnog pokušaja (do tri puta) s eksponencijalnim odmakom. Uspjeh se prepoznaje po "
 "+CMGS: odgovoru modula:"
)
code(lines("firmware/node/node.ino", 343, 364),
     caption="Listing 4.5 — sendSMS(): slanje s ponovnim pokušajima")
pagebreak()

# =========================================================
# 5. FIRMWARE - GATEWAY
# =========================================================
h1("5. Firmware — gateway")
h2("5.1. Uloga i inicijalizacija")
para(
 "Gateway prima SMS-ove od čvorova, parsira ih, šalje mjerenja na backend HTTP "
 "endpointom i svakih 30 sekundi pita backend ima li naredbi za registrirane čvorove. "
 "Inicijalizacija učitava konfiguraciju, primjenjuje eventualni reset implementacije, "
 "učitava registar čvorova iz SPIFFS-a, spaja se na WiFi, sinkronizira sat preko NTP-a "
 "i inicijalizira GSM modul:"
)
code(lines("firmware/gateway/gateway.ino", 61, 89),
     caption="Listing 5.1 — firmware/gateway/gateway.ino: setup()")

h2("5.2. Glavna petlja gatewaya")
para(
 "Glavna petlja gatewaya nadzire WiFi vezu (i ponovno se spaja po potrebi), čita "
 "dolazne bajtove s GSM modula te periodički provjerava red naredbi i pohranjene "
 "SMS poruke:"
)
code(lines("firmware/gateway/gateway.ino", 91, 113),
     caption="Listing 5.2 — Glavna petlja gatewaya")

h2("5.3. Parsiranje SMS spremnika")
para(
 "Dolazni SMS može stići inline (zajedno s +CMT: zaglavljem) ili u zasebnom retku. "
 "Funkcija processSmsBuffer razlikuje te slučajeve i prema prefiksu poruke "
 "(CAPS:, CONF: ili podatkovni CSV) usmjerava obradu odgovarajućoj funkciji:"
)
code(lines("firmware/gateway/gateway.ino", 253, 316),
     caption="Listing 5.3 — processSmsBuffer(): usmjeravanje dolaznih poruka")

h2("5.4. Obrada podatkovnog SMS-a i HTTP slanje")
para(
 "Kada stigne podatkovni SMS, gateway pronalazi čvor u registru, gradi JSON tijelo s "
 "mjerenjima (preskačući senzore čiji ID nije poznat) i šalje ga na backend. Vrijednosti "
 "NaN se preskaču pri gradnji JSON-a, pa neispravna očitanja nikada ne dolaze do baze:"
)
code(lines("firmware/gateway/gateway.ino", 318, 367),
     caption="Listing 5.4 — handleNodeSMS(): pretvaranje SMS-a u HTTP zahtjev")
para("Pomoćna funkcija addMeasurement upravo izvodi spomenuto preskakanje NaN vrijednosti:")
code(lines("firmware/gateway/gateway.ino", 653, 667),
     caption="Listing 5.5 — addMeasurement() i getIso(): gradnja JSON elemenata")

h2("5.5. Perzistentni registar čvorova")
para(
 "Registar čvorova drži se u memoriji kao polje struktura NodeRecord, ali se "
 "serijalizira u SPIFFS datoteku /nodes.json kako bi preživio restart. Time gateway "
 "nakon ponovnog pokretanja zna ID-eve čvorova i senzora bez potrebe za ponovnom "
 "registracijom:"
)
code(lines("firmware/gateway/gateway.ino", 231, 251),
     caption="Listing 5.6 — saveRegistry(): serijalizacija registra u SPIFFS")

h2("5.6. HTTP sloj s vremenskim ograničenjem")
para(
 "Sve HTTP komunikacije prolaze kroz jedinstvenu funkciju httpRequest koja postavlja "
 "API ključ u zaglavlje i — ključno — ograničava trajanje TCP veze. Bez tog ograničenja "
 "nedostupan backend bi blokirao uređaj tridesetak sekundi i okinuo watchdog:"
)
code(lines("firmware/gateway/gateway.ino", 593, 628),
     caption="Listing 5.7 — httpRequest(): HTTP sloj s timeoutom i ponovnim pokušajima")
pagebreak()

# =========================================================
# 6. SMS LIBRARY
# =========================================================
h1("6. Zajednička SMS biblioteka")
h2("6.1. Svrha i prenosivost")
para(
 "Cjelokupno parsiranje i gradnja SMS poruka izolirano je u zasebnoj C++ biblioteci "
 "(firmware/lib/sms_format) koja ne ovisi o Arduino.h zaglavlju. Zahvaljujući tome "
 "biblioteka se kompajlira pod bilo kojim standardnim C++ alatom i može se testirati "
 "na osobnom računalu (PlatformIO native cilj) bez ikakvog hardvera. To je presudno "
 "jer je upravo SMS format granica protokola između firmwarea i backenda."
)
para("Sučelje biblioteke definira strukture poruka i devet funkcija:")
code(lines("firmware/lib/sms_format/sms_format.h", 11, 42),
     caption="Listing 6.1 — sms_format.h: strukture i deklaracije funkcija")

h2("6.2. Parsiranje podatkovne poruke")
para(
 "Podatkovni SMS je CSV oblika SERIAL,temp,pressure,humidity,uv,lux nakon kojeg mogu "
 "slijediti CONF:cmdId tokeni. Funkcija parseSmsBody razlaže poruku, pretvara polja u "
 "brojeve (ili NaN) te prikuplja eventualne potvrde naredbi:"
)
code(lines("firmware/lib/sms_format/sms_format.cpp", 17, 52),
     caption="Listing 6.2 — parseSmsBody(): razlaganje podatkovnog SMS-a")
para(
 "Funkcije parseOrNan i fmtF čine simetričan par za pretvorbu između tekstualnog i "
 "brojčanog oblika, gdje sentinel vrijednost 'nan' označava neuspjelo očitanje:"
)
code(lines("firmware/lib/sms_format/sms_format.cpp", 7, 15),
     caption="Listing 6.3 — parseOrNan() i fmtF(): sentinel NaN vrijednost")

h2("6.3. Gradnja podatkovne poruke")
para(
 "Suprotan smjer — gradnja podatkovnog SMS-a iz brojčanih vrijednosti — obavlja "
 "buildNodeSms. Funkcija formatira svako polje na zadani broj decimala i prikvači "
 "potvrde naredbi na kraj poruke:"
)
code(lines("firmware/lib/sms_format/sms_format.cpp", 111, 130),
     caption="Listing 6.4 — buildNodeSms(): gradnja podatkovnog SMS-a")

h2("6.4. Naredbe i grupiranje")
para(
 "Naredbe se prenose u obliku CMD:type:payload:cmdId, a više naredbi grupira se "
 "razdvajanjem znakom točka-zarez. Funkcije parseCommand i splitBatch razlažu takvu "
 "poruku, dok buildBatchSms gradi grupiranu poruku na gateway strani:"
)
code(lines("firmware/lib/sms_format/sms_format.cpp", 54, 109),
     caption="Listing 6.5 — parseCommand(), splitBatch(), buildBatchSms()")

h2("6.5. Oglašavanje sposobnosti (CAPS)")
para(
 "Pri prvoj registraciji čvor šalje CAPS:SERIAL:SENSOR1,SENSOR2,... poruku koja "
 "navodi samo one senzore koji su uspješno inicijalizirani. Funkcije parseCapsBody i "
 "buildCapsSms obrađuju taj format:"
)
code(lines("firmware/lib/sms_format/sms_format.cpp", 132, 173),
     caption="Listing 6.6 — parseCapsBody() i buildCapsSms()")
pagebreak()

# =========================================================
# 7. BACKEND
# =========================================================
h1("7. Backend")
h2("7.1. Tehnologije")
para(
 "Backend je izgrađen na Spring Bootu 3 uz MySQL bazu, Flyway za upravljanje shemom, "
 "JWT i API-ključ autentikaciju te OpenAPI dokumentaciju renderiranu kroz Scalar "
 "sučelje na /scalar.html. Arhitektura je strogo četveroslojna: kontroleri pozivaju "
 "aplikacijske servise, aplikacijski servisi orkestriraju domenske servise, a domenski "
 "servisi pristupaju repozitorijima. Nijedan sloj ne preskače drugi."
)

h2("7.2. Slojevita arhitektura")
para("Raspored paketa odražava slojeve arhitekture:")
code(
 "controller/           HTTP sloj, @PreAuthorize zastite\n"
 "service/application/  orkestracija (visedomenske operacije, ingestija, alarmi)\n"
 "service/domain/       CRUD logika jednog entiteta\n"
 "repository/           Spring Data JPA\n"
 "model/entity/         JPA entiteti\n"
 "model/dto/            request/response DTO objekti\n"
 "model/enums/          Role, Status, Metric, SensorType, NodeCommandType...\n"
 "mapper/               entitet <-> DTO (MapStruct)\n"
 "config/               SecurityConfig, JwtAuthFilter, ApiKeyAuthFilter\n"
 "scheduler/            NodeStatusScheduler, NodeCommandExpiryScheduler\n"
 "exception/            GlobalExceptionHandler, tipizirane iznimke",
 caption="Slika 7.1 — Struktura paketa backenda"
)

h2("7.3. Tanki kontroleri")
para(
 "Kontroleri su namjerno tanki — parsiraju HTTP zahtjev, pozovu jedan aplikacijski "
 "servis i vrate odgovor. Svaki endpoint ima SpringDoc/OpenAPI anotacije. Primjer "
 "gateway kontrolera koji opslužuje sva četiri endpointa namijenjena ESP32 firmwareu:"
)
code(lines("backend/src/main/java/com/example/meteonode/controller/GatewayController.java", 38, 81),
     caption="Listing 7.1 — GatewayController: četiri gateway endpointa")

h2("7.4. Aplikacijski servisi")
para(
 "Aplikacijski servisi orkestriraju rad: drže @PreAuthorize anotacije, koordiniraju "
 "više domenskih servisa i vlasnici su transakcijskih granica. Primjer ingestije "
 "mjerenja koji za svako očitanje provjerava pripadnost senzora čvoru, sprema mjerenje "
 "i odmah pokreće evaluaciju alarma:"
)
code(lines("backend/src/main/java/com/example/meteonode/service/application/MeasurementIngestionService.java", 11, 33),
     caption="Listing 7.2 — MeasurementIngestionService: orkestracija ingestije")

h2("7.5. Centralizirana obrada pogrešaka")
para(
 "Sve iznimke obrađuje GlobalExceptionHandler koji ih mapira u odgovarajuće HTTP "
 "statuse: ResourceNotFoundException → 404, ConflictException → 409, "
 "BadRequestException → 400, AccessDeniedException → 403, a sve ostalo → 500. Time se "
 "obrada pogrešaka drži na jednom mjestu, a kontroleri ostaju čisti."
)

h2("7.6. Pregled endpointa")
para("Sustav izlaže sljedeće skupine endpointa (izbor):")
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text="Kontroler"; hdr[1].text="Bazna putanja"; hdr[2].text="Auth"; hdr[3].text="Primjeri"
for r in [
 ("AuthController","/api/auth","Javno","register, login, refresh, logout"),
 ("WeatherStationController","/api/stations","JWT","GET /, overview, status toggle"),
 ("NodeController","/api/nodes","JWT","lista, dodjela postaji, status"),
 ("MeasurementController","/api/measurements","JWT","latest, history (paginirano)"),
 ("AlarmRuleController","/api/alarm-rules","ADMIN/OPERATOR","CRUD, toggle"),
 ("GatewayController","/api/gateway","API ključ","register, measurements, commands"),
]:
    c = tbl.add_row().cells
    c[0].text,c[1].text,c[2].text,c[3].text = r
pagebreak()

# =========================================================
# 8. BAZA PODATAKA
# =========================================================
h1("8. Baza podataka")
h2("8.1. Pregled i hijerarhija")
para(
 "Shema se sastoji od jedanaest tablica kojima upravlja Flyway. Hijerarhija entiteta "
 "od vrha prema dolje izgleda ovako:"
)
code(
 "User (korisnik)\n"
 " +- WeatherStation (postaja, vlasnik je User)\n"
 "     +- Node (ESP32 cvor, opcijski vezan uz postaju)\n"
 "         +- Sensor (jedan senzor: BMP280, DHT11 ...)\n"
 "             +- Measurement (jedno ocitanje: vrijednost + metrika + vrijeme)\n"
 "             +- AlarmRule (prag koji pali alarm)\n"
 "                 +- AlarmNotification (okinuti alarm)\n"
 "                     +- AlarmNotificationRead (tko ga je procitao)\n"
 "Node\n"
 " +- NodeCommand (naredba cvoru: REBOOT, SET_INTERVAL ...)\n"
 "Gateway (zasebna tablica, API kljuc)\n"
 "RefreshToken (JWT refresh tokeni)",
 caption="Slika 8.1 — Hijerarhija tablica"
)
para(
 "Čvor ne mora biti dodijeljen postaji — kada se gateway prvi put registrira "
 "nepoznatog čvora, čvor postoji u bazi bez postaje sve dok ga administrator ili "
 "operater ne dodijeli. Stupac nodes.station_id stoga je nullable."
)

h2("8.2. Inicijalna shema")
para(
 "Prva Flyway migracija stvara sve tablice, ograničenja, strane ključeve i indekse. "
 "Izdvojeni dio koji prikazuje tablice mjerenja i alarmnih pravila:"
)
code(lines("backend/src/main/resources/db/migration/V1__schema.sql", 63, 89),
     caption="Listing 8.1 — V1__schema.sql: tablice measurements i alarm_rules")
para("Pažljivo postavljeni indeksi ubrzavaju ključne upite — povijesne grafove, "
     "evaluaciju alarma i poll naredbi:")
code(lines("backend/src/main/resources/db/migration/V1__schema.sql", 128, 142),
     caption="Listing 8.2 — Indeksi za česte upite")

h2("8.3. Evolucija sheme kroz Flyway")
para(
 "Sve promjene sheme idu isključivo kroz Flyway migracije; u Docker okruženju koristi "
 "se ddl-auto=validate, nikada create ili update. Time je povijest sheme u potpunosti "
 "rekonstruirana i ponovljiva. Neke od kasnijih migracija:"
)
bullets([
 "V4 — uvodi tablicu gateways, čini nodes.station_id nullable i uklanja api_key iz postaja.",
 "V8 — preimenuje zastarjele WAKE naredbe u REQUEST_READINGS i briše SLEEP naredbe.",
 "V9 — briše pogrešno umetnute DHT11 temperaturne zapise (BMP280 je jedini izvor temperature).",
 "V10 — dodaje last_sent_at stupac za leasing isporuke naredbi.",
])

h2("8.4. Pregled entiteta")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells; hdr[0].text="Entitet"; hdr[1].text="Ključna polja"
for r in [
 ("User","username, email, password (BCrypt), role"),
 ("WeatherStation","name, location, status, owner"),
 ("Node","serialNumber, displayName, status, lastSeen, reportingIntervalSeconds, station?"),
 ("Sensor","sensorType, isActive, node"),
 ("Measurement","value, metric, measuredAt, sensor"),
 ("NodeCommand","type, status, payload, node"),
 ("AlarmRule","metric, minValue, maxValue, severity, cooldownSeconds, isActive, sensor"),
 ("AlarmNotification","triggeredAt, value, rule"),
]:
    c = tbl.add_row().cells; c[0].text, c[1].text = r
pagebreak()

# =========================================================
# 9. AUTH
# =========================================================
h1("9. Autentikacija i autorizacija")
h2("9.1. Dvije vrste autentikacije")
para(
 "Sustav koristi dvije odvojene vrste autentikacije koje dijele isti filter lanac. "
 "Za korisnička sučelja koristi se JWT: prijava vraća 15-minutni pristupni token i "
 "7-dnevni refresh token. Za gateway se koristi API ključ pohranjen u tablici "
 "gateways, neovisan o bilo kojoj postaji."
)
para("Sigurnosna konfiguracija definira javne putanje, gateway putanje (samo API ključ) "
     "i redoslijed filtara — ApiKeyAuthFilter prije JwtAuthFilter:")
code(lines("backend/src/main/java/com/example/meteonode/config/SecurityConfig.java", 33, 53),
     caption="Listing 9.1 — SecurityConfig: filter lanac i pravila pristupa")

h2("9.2. API-ključ filter za gateway")
para(
 "ApiKeyAuthenticationFilter aktivira se samo za /api/gateway/** putanje. Traži ključ "
 "u X-Api-Key zaglavlju, pronalazi pripadni gateway i postavlja ROLE_GATEWAY ovlast. "
 "Ako gateway pošalje i svoj URL u X-Gateway-Url zaglavlju, filter ažurira zapis:"
)
code(lines("backend/src/main/java/com/example/meteonode/config/ApiKeyAuthenticationFilter.java", 30, 61),
     caption="Listing 9.2 — ApiKeyAuthenticationFilter: autentikacija gatewaya")

h2("9.3. Transparentno osvježavanje JWT tokena")
para(
 "Frontend pohranjuje tokene u localStorage i transparentno ponavlja zahtjev nakon "
 "osvježavanja. Axios interceptor presreće 401 odgovore, jednom osvježi pristupni "
 "token (s dijeljenim obećanjem da se izbjegne istovremeno višestruko osvježavanje) "
 "i ponovi izvorni zahtjev:"
)
code(lines("frontend/src/services/api.ts", 11, 45),
     caption="Listing 9.3 — frontend/src/services/api.ts: Axios refresh interceptor")

h2("9.4. Tri uloge")
para(
 "Sustav poznaje tri uloge. ADMIN ima potpun pristup. OPERATOR može izdavati naredbe, "
 "upravljati čvorovima i senzorima te alarmnim pravilima. USER ima pristup samo za "
 "čitanje. @PreAuthorize anotacije nalaze se na metodama aplikacijskih servisa, a rute "
 "frontenda zrcale ta pravila."
)
pagebreak()

# =========================================================
# 10. PIPELINE: MEASUREMENT
# =========================================================
h1("10. Procesni tok mjerenja")
para(
 "Tok mjerenja je primarni put podataka kroz sustav — od fizičkog senzora do retka u "
 "bazi. Ovo poglavlje prati taj put korak po korak, naglašavajući mjesta na kojima se "
 "podaci transformiraju ili odbacuju."
)
h2("10.1. Korak 1 — očitanje i gradnja SMS-a")
para(
 "Čvor se budi, čita sve senzore i gradi jednu SMS poruku. Neuspjela ili besmislena "
 "očitanja serijaliziraju se kao 'nan'. Tipična podatkovna poruka izgleda ovako:"
)
code("ESP32-NODE-001,24.59,1012.01,50.0,0.00,342.5",
     caption="Slika 10.1 — Primjer podatkovnog SMS-a")
para("Polja su redom: serijski broj, temperatura, tlak, vlažnost, UV indeks, osvjetljenje.")

h2("10.2. Korak 2 — prijem i parsiranje na gatewayu")
para(
 "Gateway prima SMS, prepoznaje da nije CAPS: ni CONF: poruka te ga prosljeđuje "
 "funkciji handleNodeSMS. Ona parsira CSV, pronalazi čvor u registru i gradi JSON "
 "tijelo preskačući NaN vrijednosti (vidi listing 5.4). Ako su sva očitanja NaN, "
 "HTTP zahtjev se uopće ne šalje."
)

h2("10.3. Korak 3 — HTTP ingestija na backendu")
para(
 "Gateway šalje POST /api/gateway/measurements. Aplikacijski servis "
 "MeasurementIngestionService ažurira lastSeen čvora, za svako očitanje provjerava "
 "pripada li senzor tom čvoru, sprema mjerenje i pokreće evaluaciju alarma (listing "
 "7.2). Cijela operacija je transakcijska."
)

h2("10.4. Korak 4 — povratak prema korisniku")
para(
 "Spremljena mjerenja postaju dostupna kroz MeasurementController: najnovija "
 "vrijednost (latest), sve metrike senzora (latest/all) i povijest s paginacijom i "
 "filtriranjem po vremenskom rasponu (history). Frontend te podatke prikazuje na "
 "povijesnom grafu (Recharts)."
)

h2("10.5. Ključne projektne odluke")
bullets([
 "Sentinel NaN — neispravna očitanja prenose se kao 'nan' i odbacuju na gatewayu, pa "
 "nikada ne onečišćuju bazu.",
 "Idempotentni lastSeen — svaki podatkovni SMS prirodno ažurira lastSeen, pa nije "
 "potreban poseban heartbeat.",
 "Sinkrona evaluacija alarma — namjerno odabrana na demo skali; mjerenja se uvijek "
 "spremaju, čak i ako čvor nema dodijeljenu postaju (tada se alarmi tiho preskaču).",
])
pagebreak()

# =========================================================
# 11. PIPELINE: COMMANDS
# =========================================================
h1("11. Procesni tok naredbi")
para(
 "Red naredbi (tablica node_commands) omogućuje poslužitelju da utječe na čvor unatoč "
 "tome što čvor nema stalnu vezu. Naredba prolazi kroz nekoliko statusa: PENDING → "
 "DELIVERED, uz terminalne EXPIRED i CANCELLED."
)
h2("11.1. Izdavanje naredbe")
para(
 "Naredbe REQUEST_READINGS i REBOOT izdaju se kroz generički endpoint, dok SET_INTERVAL "
 "ima namjenski endpoint. Pokušaj izdavanja SET_INTERVAL kroz generički put vraća 400. "
 "Aplikacijski servis koji to provodi:"
)
code(lines("backend/src/main/java/com/example/meteonode/service/application/NodeCommandManagementService.java", 27, 47),
     caption="Listing 11.1 — NodeCommandManagementService: izdavanje i dohvat naredbi")

h2("11.2. Isporuka putem lease/visibility prozora")
para(
 "Gateway poll-a pending naredbe svakih 30 sekundi. Naredba ostaje PENDING dok njena "
 "potvrda ne završi puni SMS krug, što redovito traje dulje od intervala poll-a. Da se "
 "ista naredba ne bi ponovno slala na svakom poll-u dok je potvrda u tijeku, svaka "
 "isporučena naredba se 'leasa' — označi vremenom last_sent_at i izuzima iz idućih "
 "poll-ova dok ne istekne prozor (zadano 180 s). Ako potvrda ne stigne, lease istječe "
 "i naredba se ponovno šalje — isporuka barem jednom (engl. at-least-once)."
)
para("Upit koji vraća isporučive naredbe poštujući lease prozor:")
code(lines("backend/src/main/java/com/example/meteonode/repository/NodeCommandRepository.java", 18, 27),
     caption="Listing 11.2 — NodeCommandRepository: findDeliverable s resend prozorom")

h2("11.3. Grupiranje i slanje prema čvoru")
para(
 "Gateway pakira do tri naredbe u jedan SMS. Za svaki čvor poll-a backend, parsira "
 "JSON listu naredbi, gradi grupiranu poruku i šalje je SMS-om:"
)
code(lines("firmware/gateway/gateway.ino", 516, 559),
     caption="Listing 11.3 — checkPendingCommands(): poll i slanje naredbi")

h2("11.4. Idempotentna potvrda i nuspojave")
para(
 "Potvrda već isporučene naredbe je no-op koji vraća uspjeh — nuspojave se izvršavaju "
 "samo pri prvoj isporuci, pa kasna duplicirana potvrda ne može npr. vratiti oporavljeni "
 "čvor u stanje OFFLINE. Nuspojave po tipu: SET_INTERVAL ažurira interval čvora, REBOOT "
 "označava čvor OFFLINE:"
)
code(lines("backend/src/main/java/com/example/meteonode/service/application/NodeCommandManagementService.java", 49, 64),
     caption="Listing 11.4 — confirmDelivered(): idempotentna potvrda i nuspojave")

h2("11.5. Dvostruka zaštita od duplikata")
para(
 "Kao druga linija obrane, firmware čvora deduplicira identifikatore naredbi obrađene "
 "unutar zadnjih 120 sekundi (listing 4.3), pa SMS koji GSM mreža ponovno isporuči "
 "nikada nije izvršen ni potvrđen dvaput. Naredbe koje ostanu PENDING 24 sata označava "
 "kao EXPIRED zaseban scheduler koji se izvršava svakih sat vremena."
)
pagebreak()

# =========================================================
# 12. PIPELINE: ALARMS
# =========================================================
h1("12. Procesni tok alarma")
para(
 "Alarmna pravila definiraju brojčane pragove na pojedinom paru senzor+metrika. Svako "
 "pravilo ima opcijske min/max pragove, ozbiljnost (INFO, WARNING, CRITICAL) i cooldown "
 "u sekundama. Evaluacija je sinkrona s ingestijom mjerenja."
)
h2("12.1. Evaluacija pri ingestiji")
para(
 "Za svako spremljeno mjerenje AlarmEvaluationService učitava aktivna pravila za taj "
 "par senzor+metrika i provjerava granice. Ako pravilo nije prekršeno ili je u "
 "cooldownu, preskače se; inače se stvara notifikacija za svakog korisnika:"
)
code(lines("backend/src/main/java/com/example/meteonode/service/application/AlarmEvaluationService.java", 27, 73),
     caption="Listing 12.1 — AlarmEvaluationService: evaluacija pravila")

h2("12.2. Atomsko paljenje uz cooldown")
para(
 "Ključni dio alarmnog toka jest sprječavanje poplave notifikacija. Provjera cooldowna "
 "i postavljanje lastFiredAt izvode se jednim atomskim SQL UPDATE-om koji uspijeva samo "
 "ako je cooldown protekao — čime se eliminira utrka (engl. race condition) između "
 "istodobnih evaluacija:"
)
code(lines("backend/src/main/java/com/example/meteonode/service/domain/AlarmRuleService.java", 84, 90),
     caption="Listing 12.2 — AlarmRuleService.tryFire(): atomsko paljenje")
para("Sam upit koji čini operaciju atomskom:")
code(lines("backend/src/main/java/com/example/meteonode/repository/AlarmRuleRepository.java", 19, 22),
     caption="Listing 12.3 — fireIfReady(): uvjetni UPDATE")
para(
 "Metoda tryFire koristi Propagation.REQUIRES_NEW kako bi se ažuriranje lastFiredAt "
 "izvršilo u zasebnoj transakciji, neovisno o ishodu ostatka evaluacije. Tek ako "
 "UPDATE uspije (vrati broj veći od nule), stvaraju se notifikacije."
)

h2("12.3. Stanje pročitanosti po korisniku")
para(
 "Stanje pročitanosti notifikacije vodi se po korisniku kroz tablicu "
 "alarm_notification_reads s kompozitnim primarnim ključem (notification_id, user_id). "
 "Označavanje pročitanim je idempotentno — duplicirani pozivi vraćaju se rano bez "
 "pristupa bazi, čime se izbjegava kršenje kompozitnog ključa."
)

h2("12.4. Rubni slučajevi")
bullets([
 "Ako senzorov čvor nema dodijeljenu postaju, evaluacija alarma se tiho preskače, ali "
 "se mjerenje uvijek sprema.",
 "Ako u sustavu nema korisnika kojima bi se poslala notifikacija, evaluacija se "
 "prekida uz upozorenje u logu.",
 "Polje AlarmNotification.read namjerno nije mapirano MapStructom — postavlja se po "
 "korisniku u servisnom sloju.",
])
pagebreak()

# =========================================================
# 13. DINAMICKA REGISTRACIJA
# =========================================================
h1("13. Dinamička registracija senzora")
para(
 "Jedna od naprednijih značajki sustava jest sposobnost čvora da se sam registrira i "
 "oglasi koje senzore stvarno posjeduje. To uklanja potrebu za ručnim konfiguriranjem "
 "svakog čvora na poslužitelju."
)
h2("13.1. Oglašavanje pri prvom pokretanju")
para(
 "Pri prvom pokretanju (ili nakon promjene serijskog broja/implementacije) čvor šalje "
 "CAPS: poruku koja navodi samo senzore koji su se uspješno inicijalizirali. NVS "
 "zastavica capsSent sprječava ponavljanje. Dio setup() funkcije zadužen za to:"
)
code(lines("firmware/node/node.ino", 162, 177),
     caption="Listing 13.1 — Oglašavanje sposobnosti pri prvom pokretanju")

h2("13.2. Usmjeravanje CAPS poruke na gatewayu")
para(
 "Gateway prepoznaje CAPS: prefiks i prosljeđuje poruku funkciji handleCapsSms, koja "
 "parsira listu senzora i registrira čvor s oglašenim senzorima:"
)
code(lines("firmware/gateway/gateway.ino", 580, 591),
     caption="Listing 13.2 — handleCapsSms(): obrada oglašavanja")
para(
 "Funkcija findOrRegisterNode zatim poziva backend (POST /nodes/register), prima "
 "dodijeljene ID-eve čvora i senzora te ih trajno pohranjuje u SPIFFS registar. Od tog "
 "trenutka gateway zna koji ID koristiti u HTTP zahtjevima za to mjerenje:"
)
code(lines("firmware/gateway/gateway.ino", 375, 433),
     caption="Listing 13.3 — findOrRegisterNode(): registracija na backendu")

h2("13.3. Rezervni mehanizam")
para(
 "Čvorovi bez prethodnog CAPS oglašavanja vraćaju se na fiksnu listu od četiri senzora. "
 "Time je osigurana kompatibilnost sa starijim čvorovima koji ne podržavaju dinamičko "
 "oglašavanje."
)
pagebreak()

# =========================================================
# 14. FRONTEND
# =========================================================
h1("14. Frontend")
h2("14.1. Tehnologije")
para(
 "Frontend je izgrađen na React + Vite + TypeScript stogu uz TailwindCSS v3, React "
 "Query v5 za dohvat podataka, React Router v6 za navigaciju, Axios za HTTP i Recharts "
 "za grafove. U Docker okruženju poslužuje ga nginx."
)
h2("14.2. Stranice")
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells; hdr[0].text="Stranica"; hdr[1].text="Tko vidi"; hdr[2].text="Što radi"
for r in [
 ("Dashboard","svi","kartice postaja, statistike"),
 ("Station detail","svi","čvorovi postaje, grafovi, dodjela čvorova"),
 ("Notifications","svi","lista alarma, označavanje pročitanog"),
 ("Alarm rules","ADMIN/OPERATOR","postavljanje pragova po senzoru i metrici"),
 ("Device management","ADMIN/OPERATOR","popis čvorova, senzora, slanje naredbi"),
 ("Admin","ADMIN","upravljanje korisnicima, sistemski pregled"),
]:
    c = tbl.add_row().cells; c[0].text,c[1].text,c[2].text = r

h2("14.3. Komunikacija s backendom")
para(
 "Sav HTTP promet prolazi kroz centraliziranu Axios instancu (listing 9.3) koja "
 "automatski dodaje JWT token i transparentno ga osvježava. React Query upravlja "
 "predmemoriranjem i ponovnim dohvatom podataka, čime se sučelje drži usklađenim s "
 "poslužiteljem bez ručnog upravljanja stanjem."
)
h2("14.4. Javni prikaz statistike")
para(
 "Prijavna stranica sadrži javnu traku koja dohvaća /api/info i prikazuje broj postaja "
 "i čvorova u stvarnom vremenu. Taj endpoint je jedini (uz autentikaciju) javno "
 "dostupan i agregira brojače kroz SystemInfoService."
)
pagebreak()

# =========================================================
# 15. TESTIRANJE
# =========================================================
h1("15. Testiranje")
h2("15.1. Strategija")
para(
 "Sustav je pokriven na tri razine: jedinični testovi poslovne logike, integracijski "
 "testovi cijelog HTTP sloja i jedinični testovi firmware protokola. Integracijski "
 "testovi koriste Testcontainers koji podiže stvarnu MySQL 8.4 bazu u kontejneru, dok "
 "jedinični i firmware testovi rade bez Dockera."
)
h2("15.2. Jedinični testovi backenda")
para(
 "AlarmEvaluationServiceTest pokriva sedam scenarija: nepostojanje aktivnih pravila "
 "preskače evaluaciju, prekoračenje gornje i donje granice okida notifikaciju, "
 "vrijednost unutar raspona ne okida, neuspjeli tryFire potiskuje notifikaciju, te "
 "rubni slučajevi s primateljima notifikacija."
)
h2("15.3. Integracijski testovi")
bullets([
 "AuthIntegrationTest — registracija, prijava, JWT pristup, 401/403/409 slučajevi, refresh.",
 "NodeProvisioningIntegrationTest — prva registracija (201), idempotentna re-registracija.",
 "MeasurementIngestionIntegrationTest — valjani batch (204), okidanje alarma, nepoznati čvor.",
 "NodeCommandIntegrationTest — izdavanje, poll, potvrda, istek i otkazivanje naredbi.",
 "AlarmRuleIntegrationTest — CRUD, toggle, potiskivanje cooldownom od kraja do kraja.",
])
h2("15.4. Firmware testovi")
para(
 "Firmware testovi (firmware/test/test_sms_format.cpp) sadrže 27 Unity testova koji "
 "pokrivaju svih devet funkcija SMS biblioteke. Pokreću se naredbom 'pio test -e "
 "native' na osobnom računalu, bez ikakvog hardvera — što je izravna posljedica "
 "izolacije SMS formata u biblioteku bez Arduino ovisnosti (poglavlje 6)."
)
pagebreak()

# =========================================================
# 16. POKRETANJE
# =========================================================
h1("16. Pokretanje i implementacija")
h2("16.1. Pokretanje putem Dockera")
para("Cijeli sustav podiže se jednom naredbom uz Docker:")
code(
 "# Sve putem Dockera\n"
 "docker compose up --build\n"
 "# Frontend  -> http://localhost:3000\n"
 "# API + docs -> http://localhost:8080/scalar.html\n\n"
 "# Samo backend (lokalni MySQL na portu 3308)\n"
 "./mvnw spring-boot:run\n\n"
 "# Testovi\n"
 "./mvnw test\n\n"
 "# Firmware jedinicni testovi (bez hardvera)\n"
 "cd firmware && pio test -e native",
 caption="Listing 16.1 — Naredbe za pokretanje sustava"
)
h2("16.2. Početni korisnici")
para("Inicijalne migracije sjeme sustava s tri korisnika (lozinka admin123 za sve):")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells; hdr[0].text="Korisničko ime"; hdr[1].text="Uloga"
for r in [("admin","ADMIN"),("operator","OPERATOR"),("viewer","USER")]:
    c = tbl.add_row().cells; c[0].text,c[1].text = r
h2("16.3. Važna ograničenja implementacije")
bullets([
 "Slojevita arhitektura je fiksna — nijedan sloj ne preskače drugi.",
 "Imena enumeracija su format protokola — SensorType i NodeCommandType vrijednosti su "
 "točni nizovi koje firmware šalje; preimenovanje zahtijeva Flyway migraciju, flešanje "
 "firmwarea i ponovnu registraciju čvorova.",
 "Promjene sheme idu isključivo kroz Flyway; nikada ddl-auto=create ili update u Dockeru.",
 "Ugovor gateway API-ja je zamrznut — četiri gateway endpointa ne mogu mijenjati oblik "
 "bez koordiniranog ažuriranja firmwarea.",
])
pagebreak()

# =========================================================
# 17. ZAKLJUČAK
# =========================================================
h1("17. Zaključak")
para(
 "MeteoNode pokazuje da se cjelovita platforma za nadzor okoliša može izgraditi i oko "
 "naizgled ograničavajuće pretpostavke da senzorski čvorovi nemaju pristup internetu. "
 "Pretvaranjem SMS-a u jedini komunikacijski kanal čvora, sustav postaje primjenjiv na "
 "udaljenim lokacijama na kojima je GSM signal jedina dostupna infrastruktura."
)
para(
 "Glavna pouka projekta jest važnost jasno definirane granice protokola. SMS format "
 "poruka predstavlja ugovor između triju neovisno isporučivih artefakata, a njegova "
 "izolacija u biblioteku bez Arduino ovisnosti omogućila je testiranje protokola na "
 "računalu i dala čvrst temelj cijelom sustavu. Tri analizirana procesna toka — "
 "mjerenje, naredbe i alarmi — pokazuju kako se kroz tu granicu provlače podaci, "
 "upravljanje i reakcije sustava, svaki sa svojim mehanizmima otpornosti: sentinel "
 "NaN vrijednostima, lease prozorom za isporuku barem jednom te atomskim paljenjem "
 "alarma uz cooldown."
)
para(
 "Sustav je u potpunosti implementiran i testiran na trima razinama. Mogući smjerovi "
 "daljnjeg razvoja uključuju proširenje broja podržanih čvorova po gatewayu, asinkronu "
 "evaluaciju alarma na većoj skali te bogatiju vizualizaciju povijesnih podataka na "
 "nadzornoj ploči."
)

h1("Literatura")
para("[1] Spring Boot Reference Documentation, https://docs.spring.io/spring-boot/")
para("[2] Espressif ESP32 Technical Reference Manual.")
para("[3] SIMCom SIM800L AT Command Manual.")
para("[4] React Documentation, https://react.dev/")
para("[5] Flyway Documentation, https://flywaydb.org/documentation/")
para("[6] PlatformIO i Unity test framework, https://platformio.org/")

out = os.path.join(ROOT, "MeteoNode_zavrsni_rad.docx")
doc.save(out)
print("Saved:", out)
print("Paragraphs:", len(doc.paragraphs))
